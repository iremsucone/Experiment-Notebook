import json
import datetime
import os
from fpdf import FPDF
from docx import Document
from docx.shared import Inches

def get_input_list(prompt, detail_prompt=None):
    """
    Collect a list of user inputs until an empty string is entered.
    If a detail_prompt is provided, each item will prompt for an additional detail.
    """
    items = []
    while True:
        item = input(prompt)
        if not item:
            break
        if detail_prompt:
            detail = input(detail_prompt.format(item=item))
            item = f"{item} ({detail})"
        items.append(item)
    return items

def get_experiment_details():
    """
    Collects all experiment details from the user and returns
    the experiment dictionary along with a timestamp.
    """
    now = datetime.datetime.now()
    timestamp = now.strftime('%Y%m%d%H%M%S')
    experiment = {
        "title": input("Experiment Title: "),
        "experiment_id": f"EXP-{timestamp}",
        "date": str(now),
        "experimenter": input("Experimenter Name: "),
        "project": input("Project Name: "),
        "chemicals": get_input_list("Enter chemical (or press Enter to finish): ",
                                     "Amount of {item} used: "),
        "equipment": get_input_list("Enter equipment (or press Enter to finish): "),
        "procedure": get_input_list("Enter procedure step (or press Enter to finish): "),
        "observations": input("Observations: "),
        "results": input("Results: ")
    }
    return experiment, timestamp

def load_experiments(filepath="experiment_log.json"):
    """
    Loads existing experiments from a JSON file.
    Returns an empty list if the file doesn't exist or is invalid.
    """
    if os.path.exists(filepath):
        try:
            with open(filepath, "r") as file:
                data = json.load(file)
                if isinstance(data, list):
                    return data
        except json.JSONDecodeError:
            pass
    return []

def save_experiments(experiments, filepath="experiment_log.json"):
    """
    Saves the list of experiments to a JSON file.
    """
    with open(filepath, "w") as file:
        json.dump(experiments, file, indent=4)

def generate_pdf_report(experiment, timestamp):
    """
    Generates a PDF report with table-like structures for the experiment details.
    Procedure steps are numbered.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Dimensions for table cells
    col_width_field = 50
    col_width_value = 130
    row_height = 10
    
    # Basic Information Table
    fields = [
        ("Title", experiment["title"]),
        ("Experiment ID", experiment["experiment_id"]),
        ("Date", experiment["date"]),
        ("Experimenter", experiment["experimenter"]),
        ("Project", experiment["project"])
    ]
    
    # Header row with a gray background
    pdf.set_fill_color(200, 200, 200)
    pdf.cell(col_width_field, row_height, "Field", border=1, fill=True)
    pdf.cell(col_width_value, row_height, "Value", border=1, fill=True, ln=True)
    
    pdf.set_fill_color(255, 255, 255)
    for field, value in fields:
        pdf.cell(col_width_field, row_height, field, border=1)
        pdf.cell(col_width_value, row_height, str(value), border=1, ln=True)
    
    pdf.ln(10)
    
    # Helper to add list sections as tables with optional numbering
    def add_list_table(title, items, numbered=False):
        if items:
            pdf.set_fill_color(200, 200, 200)
            pdf.cell(0, row_height, title, border=1, fill=True, ln=True)
            pdf.set_fill_color(255, 255, 255)
            for i, item in enumerate(items, start=1):
                text = f"{i}. {item}" if numbered else f"- {item}"
                pdf.cell(0, row_height, text, border=1, ln=True)
            pdf.ln(5)
    
    add_list_table("Chemicals Used", experiment.get("chemicals", []))
    add_list_table("Equipment Used", experiment.get("equipment", []))
    add_list_table("Procedure", experiment.get("procedure", []), numbered=True)
    
    # Helper to add multiline text sections
    def add_text_section(title, text):
        if text:
            pdf.set_fill_color(200, 200, 200)
            pdf.cell(0, row_height, title, border=1, fill=True, ln=True)
            pdf.set_fill_color(255, 255, 255)
            pdf.multi_cell(0, row_height, text, border=1)
            pdf.ln(5)
    
    add_text_section("Observations", experiment.get("observations", ""))
    add_text_section("Results", experiment.get("results", ""))
    
    pdf_filename = f"experiment_report_{timestamp}.pdf"
    pdf.output(pdf_filename)
    print(f"PDF report generated: {pdf_filename}")

def generate_word_report(experiment, timestamp):
    """
    Generates a Word (DOCX) report with table structures for the experiment details.
    Procedure steps are numbered.
    """
    doc = Document()
    doc.add_heading("Experiment Report", level=1)
    
    # Basic Information Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'LightShading-Accent1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    fields = [
        ("Title", experiment["title"]),
        ("Experiment ID", experiment["experiment_id"]),
        ("Date", experiment["date"]),
        ("Experimenter", experiment["experimenter"]),
        ("Project", experiment["project"])
    ]
    
    for field, value in fields:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = str(value)
    
    doc.add_paragraph()  # spacer
    
    # Helper to add list sections with optional numbering
    def add_list_section(title, items, numbered=False):
        if items:
            doc.add_heading(title, level=2)
            for i, item in enumerate(items, start=1):
                if numbered:
                    doc.add_paragraph(f"{i}. {item}", style='List Number')
                else:
                    doc.add_paragraph(f"- {item}", style='List Bullet')
    
    add_list_section("Chemicals Used", experiment.get("chemicals", []))
    add_list_section("Equipment Used", experiment.get("equipment", []))
    add_list_section("Procedure", experiment.get("procedure", []), numbered=True)
    
    # Helper to add text sections
    def add_text_section(title, text):
        if text:
            doc.add_heading(title, level=2)
            doc.add_paragraph(text)
    
    add_text_section("Observations", experiment.get("observations", ""))
    add_text_section("Results", experiment.get("results", ""))
    
    doc_filename = f"experiment_report_{timestamp}.docx"
    doc.save(doc_filename)
    print(f"Word report generated: {doc_filename}")

def log_experiment():
    """
    Orchestrates the experiment logging process:
    - Gathers user input
    - Updates the experiment log
    - Generates PDF and Word reports with table-based structures.
    """
    experiment, timestamp = get_experiment_details()
    experiments = load_experiments()
    experiments.append(experiment)
    save_experiments(experiments)
    print("\nExperiment logged successfully! Generating reports...")
    generate_pdf_report(experiment, timestamp)
    generate_word_report(experiment, timestamp)

def main():
    log_experiment()

if __name__ == "__main__":
    main()
